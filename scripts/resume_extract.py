"""
Extracts structured, research-relevant information from a resume (PDF or
DOCX) using Claude. Output schema is designed to line up with the Purdue
faculty research-interest data (data/purdue_cs_faculty.json) so the two
can later be matched by an agent.

Usage:
    export ANTHROPIC_API_KEY=sk-...
    python3 scripts/resume_extract.py path/to/resume.pdf [-o output.json]
"""

import argparse
import json
import sys
from pathlib import Path

import anthropic

MODEL = "claude-sonnet-5"

SCHEMA_INSTRUCTIONS = """
You are extracting structured information from a student's resume to power
a research-advisor matching tool. Read the resume text below and return
ONLY a single JSON object (no markdown fences, no commentary) with exactly
this shape:

{
  "name": string | null,
  "contact": {
    "email": string | null,
    "phone": string | null,
    "location": string | null,
    "linkedin": string | null,
    "github": string | null,
    "website": string | null
  },
  "education": [
    {
      "institution": string,
      "degree": string | null,
      "field_of_study": string | null,
      "gpa": string | null,
      "start_date": string | null,
      "end_date": string | null,
      "relevant_coursework": [string]
    }
  ],
  "research_experience": [
    {
      "title": string,
      "organization": string | null,
      "advisor_or_mentor": string | null,
      "start_date": string | null,
      "end_date": string | null,
      "description": string,
      "methods_and_tools": [string]
    }
  ],
  "work_experience": [
    {
      "title": string,
      "organization": string | null,
      "start_date": string | null,
      "end_date": string | null,
      "description": string
    }
  ],
  "projects": [
    {
      "title": string,
      "description": string,
      "tools": [string]
    }
  ],
  "publications": [string],
  "skills": {
    "programming_languages": [string],
    "technical_skills": [string],
    "tools_and_frameworks": [string]
  },
  "honors_and_awards": [string],
  "research_interest_summary": string,
  "research_keywords": [string]
}

Guidance:
- "research_interest_summary" should be a 2-4 sentence synthesis, in your
  own words, of what research areas/questions this person seems genuinely
  interested in or experienced with, based on their research experience,
  projects, coursework, and publications combined -- not just a copy of
  one resume line.
- "research_keywords" should be 5-15 short topical keywords/phrases (e.g.
  "reinforcement learning", "database indexing", "computer vision") suitable
  for matching against a professor's research-area tags and bio text.
- If a field is not present on the resume, use null (for scalars) or an
  empty list (for arrays/lists). Do not invent information.
- Return valid JSON only.

Resume text:
---
{resume_text}
---
"""


def extract_text_from_pdf(path: Path) -> str:
    import pdfplumber

    text_parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text_parts.append(page.extract_text() or "")
    return "\n".join(text_parts)


def extract_text_from_docx(path: Path) -> str:
    import docx

    document = docx.Document(path)
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(path)
    if suffix in (".docx", ".doc"):
        if suffix == ".doc":
            raise ValueError(
                "Legacy .doc format is not supported -- please convert to .docx or .pdf"
            )
        return extract_text_from_docx(path)
    raise ValueError(f"Unsupported file type: {suffix} (expected .pdf or .docx)")


def call_claude_for_structured_data(resume_text: str) -> dict:
    client = anthropic.Anthropic()
    prompt = SCHEMA_INSTRUCTIONS.replace("{resume_text}", resume_text)

    response = client.messages.create(
        model=MODEL,
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}],
    )
    # Sonnet 5 runs adaptive thinking by default, so a ThinkingBlock may
    # precede the TextBlock -- find the text block rather than assuming index 0.
    text_block = next(b for b in response.content if b.type == "text")
    raw = text_block.text.strip()

    # Defensive: strip markdown code fences if the model adds them anyway.
    if raw.startswith("```"):
        raw = raw.strip("`")
        if raw.startswith("json"):
            raw = raw[len("json"):]
        raw = raw.strip()

    return json.loads(raw)


def parse_resume(path: Path) -> dict:
    resume_text = extract_text(path)
    if not resume_text.strip():
        raise ValueError(f"No extractable text found in {path}")
    structured = call_claude_for_structured_data(resume_text)
    structured["_source_file"] = str(path)
    return structured


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("resume_path", type=Path, help="Path to a .pdf or .docx resume")
    parser.add_argument(
        "-o", "--output", type=Path, default=None,
        help="Output JSON path (default: <resume_name>.extracted.json next to input)",
    )
    args = parser.parse_args()

    if not args.resume_path.exists():
        print(f"File not found: {args.resume_path}", file=sys.stderr)
        sys.exit(1)

    output_path = args.output or args.resume_path.with_suffix(".extracted.json")

    structured = parse_resume(args.resume_path)

    with open(output_path, "w") as f:
        json.dump(structured, f, indent=2, ensure_ascii=False)

    print(f"Wrote structured resume data to {output_path}", file=sys.stderr)


if __name__ == "__main__":
    main()
